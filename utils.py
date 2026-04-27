import io
import os
import streamlit as st
import anthropic
import openai

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


def get_client():
    return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])


def get_openai_client():
    return openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])


# ---------------------------------------------------------------------------
# PDF formatting extraction
# ---------------------------------------------------------------------------

def _extract_pdf_formatting(raw: bytes) -> tuple[str, str]:
    """Parse a PDF with pdfplumber and return (plain_text, formatting_guide).

    plain_text      — full text content, line by line
    formatting_guide — structured description of fonts/sizes per content role,
                       ready to be injected into an LLM prompt.
    """
    import pdfplumber

    all_lines: list[dict] = []

    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            if not page.chars:
                continue

            # Group characters into lines by y-position (2 pt bucket tolerance)
            line_map: dict[int, list] = {}
            for char in page.chars:
                bucket = round(char["top"] / 2) * 2
                line_map.setdefault(bucket, []).append(char)

            for bucket in sorted(line_map):
                chars = line_map[bucket]
                text = "".join(c["text"] for c in chars).strip()
                if not text:
                    continue

                fontnames = [c["fontname"] for c in chars]
                sizes     = [c["size"]     for c in chars]

                raw_font = max(set(fontnames), key=fontnames.count)
                # Strip PDF subset prefix (e.g. "AAAAAA+TimesNewRomanPS-BoldMT" → "TimesNewRomanPS-BoldMT")
                dominant_font = raw_font.split("+", 1)[-1]
                avg_size      = round(sum(sizes) / len(sizes), 1)
                is_bold   = any(k in dominant_font.lower() for k in ("bold", "black", "heavy"))
                is_italic = any(k in dominant_font.lower() for k in ("italic", "oblique"))

                all_lines.append({
                    "text":   text,
                    "font":   dominant_font,
                    "size":   avg_size,
                    "bold":   is_bold,
                    "italic": is_italic,
                })

    plain_text = "\n".join(ln["text"] for ln in all_lines)

    # Build formatting guide: group by (font, size, bold, italic), sorted large→small
    seen: dict[tuple, list[str]] = {}
    for ln in all_lines:
        key = (ln["font"], ln["size"], ln["bold"], ln["italic"])
        seen.setdefault(key, []).append(ln["text"])

    guide: list[str] = [
        "TEMPLATE FORMATTING SPECIFICATION",
        "Apply these exact typographic conventions to the matching content types:\n",
    ]
    for (font, size, bold, italic), examples in sorted(seen.items(), key=lambda x: -x[0][1]):
        attrs = []
        if bold:
            attrs.append("Bold")
        if italic:
            attrs.append("Italic")
        attr_str = f", {', '.join(attrs)}" if attrs else ""
        sample = " | ".join(f'"{e[:60]}"' for e in examples[:3] if e.strip())
        guide.append(f"  • {font}{attr_str}, {size}pt\n    Used for: {sample}")

    return plain_text, "\n".join(guide)


# ---------------------------------------------------------------------------
# Template loading (repo defaults + user uploads)
# ---------------------------------------------------------------------------

def load_repo_template(name: str) -> tuple[str, str]:
    """Read a template from templates/ and return (plain_text, formatting_guide).

    Tries <name>.pdf first (richest metadata), then .docx, then .txt.
    formatting_guide is non-empty only for PDFs; empty string for other formats.
    """
    for ext in (".pdf", ".docx", ".txt"):
        path = os.path.join(_TEMPLATE_DIR, f"{name}{ext}")
        if not os.path.exists(path):
            continue
        with open(path, "rb") as f:
            raw = f.read()

        if ext == ".pdf":
            try:
                return _extract_pdf_formatting(raw)
            except Exception as e:
                st.warning(f"Could not parse PDF template formatting: {e}")
                return "", ""

        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                return text, ""
            except Exception as e:
                st.warning(f"Could not read DOCX template: {e}")
                return "", ""

        if ext == ".txt":
            return raw.decode("utf-8", errors="replace"), ""

    return "", ""


def extract_template_from_upload(uploaded_file) -> tuple[str, str]:
    """Extract (plain_text, formatting_guide) from a user-uploaded template file.

    PDF uploads get full formatting metadata; DOCX and TXT return empty guide.
    """
    raw = uploaded_file.read()
    file_type = uploaded_file.type

    if file_type == "application/pdf":
        try:
            return _extract_pdf_formatting(raw)
        except Exception as e:
            st.error(f"Could not parse uploaded PDF template: {e}")
            return "", ""

    if "wordprocessingml" in file_type:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text, ""
        except Exception as e:
            st.error(f"Could not read DOCX template: {e}")
            return "", ""

    # Plain text
    return raw.decode("utf-8", errors="replace"), ""


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def generate_resume_pdf(text: str) -> bytes:
    """Render resume text (markdown-ish LLM output) to a formatted PDF.

    Styling mirrors the repo template:
      - 17 pt Times-Bold   → candidate name  (first non-empty line)
      - 10 pt Times-Roman  → contact / tagline line (second non-empty line)
      - 12 pt Times-Bold   → section headers (ALL-CAPS short lines)
      - 10 pt Times-Bold   → company / school names  (**bold** markers)
      - 10 pt Times-Italic → job titles / dates       (*italic* markers)
      - 10 pt Times-Roman  → body text
      - 10 pt Times-Roman  → bullet points (-, •, ●)
    """
    import re
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    # ── Styles ──────────────────────────────────────────────────────────────
    def _style(name, font, size, **kw):
        return ParagraphStyle(name, fontName=font, fontSize=size,
                              leading=size * 1.25, **kw)

    S = {
        "name":    _style("name",    "Times-Bold",   17, alignment=1, spaceAfter=2),
        "contact": _style("contact", "Times-Roman",  10, alignment=1, spaceAfter=6),
        "section": _style("section", "Times-Bold",   12, spaceBefore=8, spaceAfter=3),
        "company": _style("company", "Times-Bold",   10, spaceAfter=1),
        "italic":  _style("italic",  "Times-Italic", 10, spaceAfter=1),
        "body":    _style("body",    "Times-Roman",  10, spaceAfter=2),
        "bullet":  _style("bullet",  "Times-Roman",  10, spaceAfter=1,
                          leftIndent=16, firstLineIndent=-8),
    }

    # ── Inline markdown → reportlab XML ─────────────────────────────────────
    def _inline(line: str) -> str:
        """Convert **bold** and *italic* markdown to reportlab XML tags."""
        line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        line = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
        line = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", line)
        return line

    # ── Classify each line ───────────────────────────────────────────────────
    def _classify(line: str, line_idx: int, first_idx: int, second_idx: int):
        stripped = line.strip()
        if not stripped:
            return None, None

        # Markdown heading levels
        if stripped.startswith("### "):
            return "company", stripped[4:]
        if stripped.startswith("## "):
            return "section", stripped[3:]
        if stripped.startswith("# "):
            return "name", stripped[2:]

        # Bullet markers (flexible: dash, bullet chars, optional extra spaces)
        m = re.match(r"^[-•●]\s*", stripped)
        if m:
            return "bullet", stripped[m.end():]

        # **entire line bold** → company/school
        if re.match(r"^\*\*.+\*\*$", stripped):
            return "company", stripped[2:-2]

        # *entire line italic* → title/date
        if re.match(r"^\*.+\*$", stripped) and not stripped.startswith("**"):
            return "italic", stripped[1:-1]

        # ALL-CAPS short line → section header
        if stripped.isupper() and len(stripped) < 40:
            return "section", stripped

        # Positional heuristics (no markdown headers present)
        if line_idx == first_idx:
            return "name", stripped
        if line_idx == second_idx:
            return "contact", stripped

        return "body", stripped

    # ── Build story ──────────────────────────────────────────────────────────
    lines = text.splitlines()
    non_empty = [i for i, l in enumerate(lines) if l.strip()]
    first_idx  = non_empty[0] if len(non_empty) > 0 else -1
    second_idx = non_empty[1] if len(non_empty) > 1 else -1

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch,  bottomMargin=0.75 * inch,
    )

    # State: once we pass an italic (title/dates) line inside a section,
    # plain body lines are experience/project bullets until the next section resets.
    after_title = False

    story = []
    for idx, line in enumerate(lines):
        kind, content = _classify(line, idx, first_idx, second_idx)
        if kind is None:
            story.append(Spacer(1, 4))
            continue

        # Update state
        if kind == "italic":
            after_title = True
        elif kind in ("section", "name", "contact"):
            after_title = False

        # Promote unmarked body lines to bullets when inside an experience/project block
        if kind == "body" and after_title:
            kind = "bullet"

        if kind == "section":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.black, spaceAfter=2))

        rendered = f"● {_inline(content)}" if kind == "bullet" else _inline(content)
        story.append(Paragraph(rendered, S[kind]))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Cover letter PDF generation
# ---------------------------------------------------------------------------

def generate_cover_letter_pdf(text: str) -> bytes:
    """Render cover letter text (LLM output) to a formatted PDF.

    Mirrors the repo cover letter template:
      - 20 pt Times-Roman centered → name / tagline        (first non-empty line)
      - 11 pt Times-Roman centered → contact info           (second non-empty line)
      - 11 pt Times-Roman tight    → date + recipient block (up to "Dear …")
      - 11 pt Times-Roman spaced   → salutation, body paragraphs, closing
    """
    import re
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    def _style(name, font, size, **kw):
        return ParagraphStyle(name, fontName=font, fontSize=size,
                              leading=size * 1.3, **kw)

    S = {
        "name":    _style("name",    "Times-Roman", 20, alignment=1, spaceAfter=2),
        "contact": _style("contact", "Times-Roman", 11, alignment=1, spaceAfter=6),
        # tight: date + recipient address — single-spaced, no extra gap between lines
        "tight":   _style("tight",   "Times-Roman", 11, spaceAfter=1),
        # body: letter paragraphs — space between paragraphs
        "body":    _style("body",    "Times-Roman", 11, spaceAfter=8),
    }

    def _inline(line: str) -> str:
        line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        line = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
        line = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", line)
        return line

    lines = text.splitlines()
    non_empty = [i for i, l in enumerate(lines) if l.strip()]
    first_idx  = non_empty[0] if len(non_empty) > 0 else -1
    second_idx = non_empty[1] if len(non_empty) > 1 else -1

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=0.5 * inch, bottomMargin=inch,
    )

    # States: "header" → "address" → "body"
    # "address" = date + recipient block (tight spacing, no blank-line gaps)
    # Transitions to "body" on the first "Dear …" line.
    state = "header"

    story = []
    for idx, line in enumerate(lines):
        stripped = line.strip()

        # Blank line handling per state
        if not stripped:
            if state == "address":
                pass  # suppress blank lines inside the address block
            elif state == "body":
                story.append(Spacer(1, 4))
            else:
                story.append(Spacer(1, 4))
            continue

        # Detect transition from address block to letter body
        if state == "address" and re.match(r"^Dear\b", stripped, re.IGNORECASE):
            state = "body"
            story.append(Spacer(1, 10))  # single gap before salutation

        # Assign style
        if idx == first_idx:
            style = S["name"]
        elif idx == second_idx:
            style = S["contact"]
            state = "address"           # everything after contact = address block
        elif state == "address":
            style = S["tight"]
        else:
            style = S["body"]

        story.append(Paragraph(_inline(stripped), style))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Resume content helpers
# ---------------------------------------------------------------------------

def extract_text_from_file(uploaded_file):
    """Extract plain text from a PDF, DOCX, or TXT resume upload."""
    content = uploaded_file.read()
    file_type = uploaded_file.type

    if file_type == "text/plain":
        return content.decode("utf-8")

    if file_type == "application/pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return ""

    if "wordprocessingml" in file_type:
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            st.error(f"Could not read DOCX: {e}")
            return ""

    return content.decode("utf-8", errors="replace")


def resume_inputs():
    """Render the shared resume upload + text area and return combined content.

    SHORT-TERM MEMORY  – the text area is bound to st.session_state['resume_text']
                         so it persists when the user switches between pages.
    LONG-TERM MEMORY   – on the very first script run per browser session the
                         last-saved resume is pre-loaded from SQLite via memory.py.
    """
    # --- LONG-TERM MEMORY: pre-load saved resume once per browser session ---
    import memory as _mem
    _mem.init_session_resume()

    st.subheader("Your Background")

    if st.session_state.get("_resume_from_db") and st.session_state.get("resume_text"):
        st.caption("Resume pre-loaded from your last session — edit below or upload a new file.")

    uploaded = st.file_uploader("Upload Resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

    # SHORT-TERM MEMORY: key binds to session_state so content survives page switches
    pasted = st.text_area(
        "Or paste your resume / additional context here",
        key="resume_text",
        height=250,
        placeholder="Work history, skills, education, achievements…",
    )

    parts = []
    if uploaded:
        extracted = extract_text_from_file(uploaded)
        if extracted:
            parts.append(f"[UPLOADED RESUME]\n{extracted}")
    if pasted.strip():
        parts.append(pasted.strip())

    combined = "\n\n".join(parts)

    # SHORT-TERM: cache the combined result for cross-page access
    if combined:
        st.session_state["resume_content"] = combined

    return combined
